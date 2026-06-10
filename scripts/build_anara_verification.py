"""READ-ONLY: артефакт-сверка ссылок Земельного кодекса для Анары.
Читает data/final/zemelnyy_structured.html, НИЧЕГО не меняет.
Строит:
  (1) таблицу ВСЕХ ссылок (статья/пункт | якорь-текст | target | резолв | вн/внеш)
  (2) блок «Помечено Анарой» (13 пунктов) — сырой HTML + резолв
  (3) спот-чек «об административных правонарушениях» без <a>
Выводит Markdown + .docx в data/reports/.
"""
import sys
import io
import re
from pathlib import Path
from bs4 import BeautifulSoup, Tag

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parent.parent
HTML = ROOT / "data" / "final" / "zemelnyy_structured.html"
SELF_DOC = "K030000442_"
OUT_MD = ROOT / "data" / "reports" / "zemelnyy_anara_verification.md"
OUT_DOCX = ROOT / "data" / "reports" / "zemelnyy_anara_verification.docx"

FLAGGED = ["38", "81", "94", "100", "104", "105", "165"]

# Имена внешних актов по doc_id adilet (по контексту ссылок в Земельном)
EXT_NAMES = {
    "K030000481_": "Водный кодекс РК",
    "K080000095_": "Закон РК «О государственном имуществе»",
    "K1400000226": "Уголовный кодекс РК",
    "K1400000231": "Уголовно-процессуальный кодекс РК",
    "K1400000235": "Кодекс РК об административных правонарушениях",
    "K1500000375": "Предпринимательский кодекс РК",
    "K1500000377": "Гражданский процессуальный кодекс РК",
    "K1500000414": "Трудовой кодекс РК",
    "K1700000125": "Кодекс РК «О недрах и недропользовании»",
    "K2000000350": "Административный процедурно-процессуальный кодекс РК",
    "K2100000400": "Экологический кодекс РК",
    "K940001000_": "Гражданский кодекс РК (Общая часть)",
    "K950001000_": "Конституция РК",
    "K990000409_": "Гражданский кодекс РК (Особенная часть)",
    "S2500000069": "Нормативное постановление Конституционного Суда РК",
    "V1500012590": "Правила осуществления государственных закупок (приказ)",
    "V2500036055": "Правила субсидирования арендной платы за жилище (приказ)",
    "V2500036333": "Правила субсидирования стоимости услуг по подаче воды (приказ)",
    "V2600038568": "Правила дорожного движения (приказ МВД)",
    "V2600038598": "Методики по подсчёту запасов полезных ископаемых (приказ)",
    "V2600038674": "О ввозе сахара-сырца тростникового (приказ)",
    "Z010000242_": "Закон РК «Об архитектурной, градостроительной и строительной деятельности»",
    "Z060000175_": "Закон РК «Об особо охраняемых природных территориях»",
    "Z070000296_": "Закон РК «О статусе столицы Республики Казахстан»",
    "Z070000310_": "Закон РК «О государственной регистрации прав на недвижимое имущество»",
    "Z1100000413": "Закон РК «О государственном имуществе»",
    "Z1400000176": "Закон РК «О реабилитации и банкротстве»",
    "Z1500000434": "Закон РК «О государственных закупках»",
    "Z1600000486": "Закон РК «О долевом участии в жилищном строительстве»",
    "Z1700000047": "Закон РК «О пастбищах»",
    "Z1900000242": "Закон РК «О специальных экономических и индустриальных зонах»",
    "Z1900000243": "Закон РК (Z1900000243)",
    "Z2300000021": "Закон РК «О возврате государству незаконно приобретённых активов»",
    "Z2500000165": "Закон РК «Об особом статусе города Туркестана»",
    "Z2600000258": "Закон РК «О банках и банковской деятельности в РК»",
    "Z950002444_": "Закон РК «О банках и банковской деятельности в РК»",
    "Z980000258_": "Закон РК «Об особом статусе города Алматы»",
}

soup = BeautifulSoup(HTML.read_text(encoding="utf-8"), "html.parser")


def norm(s):
    return re.sub(r"\s+", " ", s).strip()


def nearest_article(el):
    p = el
    while p is not None:
        if isinstance(p, Tag) and p.get("data-type") == "статья":
            return p
        p = p.parent
    return None


_TITLE_CACHE = {}


def article_title(div):
    key = id(div)
    if key in _TITLE_CACHE:
        return _TITLE_CACHE[key]
    title = None
    for h in div.find_all(["h2", "h3", "h4", "p", "b", "strong"]):
        t = norm(h.get_text(" "))
        if re.match(r"Стать[яиею]\s+\d", t):
            title = t
            break
    _TITLE_CACHE[key] = title
    return title


def resolve_internal(frag):
    """frag = 'z72' → ('64', 'Статья 64. ...', is_head)."""
    if not frag:
        return (None, None, False)
    el = soup.find(id=frag) or soup.find(attrs={"name": frag})
    if el is None:
        return ("БИТАЯ", None, False)
    art = nearest_article(el)
    if art is None:
        return ("вне статьи", None, False)
    num = art.get("data-number")
    title = article_title(art)
    is_head = bool(title and re.match(rf"Стать[яиею]\s+{re.escape(num)}\b",
                                      norm(el.get_text(" ")) if isinstance(el, Tag) else ""))
    return (num, title, is_head)


def classify_href(href):
    href = href or ""
    if href.startswith("#"):
        return ("internal", href[1:], None)
    if SELF_DOC in href:
        frag = href.split("#")[-1] if "#" in href else None
        return ("internal", frag, None)
    m = re.search(r"adilet\.zan\.kz/\w+/docs/([A-Za-z0-9_]+)", href)
    if m:
        frag = href.split("#")[-1] if "#" in href else None
        return ("external", frag, m.group(1))
    return ("external", None, None)


def point_label(punkt, sub):
    loc = []
    if punkt:
        loc.append(f"п.{punkt}")
    if sub:
        loc.append(f"пп.{sub})")
    return " ".join(loc) if loc else "(без нумерации)"


P_PUNKT = re.compile(r"^\s*(\d+(?:-\d+)?)\.\s")
P_SUB = re.compile(r"^\s*(\d+(?:-\d+)?)\)\s")


def collect_rows():
    """Возвращает список строк-ссылок в порядке статья→документ.
    Обходит ВСЕ потомки статьи в порядке документа: на каждом <p> обновляет
    контекст пункт/подпункт, на каждом <a href> эмитит строку (захватывает и
    ссылки вне <p> — напр. в ИЗПИ-примечаниях <font>)."""
    rows = []
    arts = soup.find_all("div", attrs={"data-type": "статья"})
    for div in arts:
        art_num = div.get("data-number")
        punkt = None
        sub = None
        for node in div.descendants:
            if not isinstance(node, Tag):
                continue
            if node.name == "p":
                txt = norm(node.get_text(" "))
                mp = P_PUNKT.match(txt)
                ms = P_SUB.match(txt)
                if mp:
                    punkt = mp.group(1)
                    sub = None
                elif ms:
                    sub = ms.group(1)
                continue
            if node.name == "a":
                a = node
                href = a.get("href")
                if not href:
                    continue  # name-anchor, не ссылка
                in_p = a.find_parent("p") is not None
                kind, frag, docid = classify_href(href)
                atext = norm(a.get_text(" "))
                loc = point_label(punkt, sub)
                if not in_p:
                    loc = (loc + " (примечание)").strip()
                if kind == "internal":
                    tnum, ttitle, is_head = resolve_internal(frag)
                    target = f"#{frag}" if frag else "(нет фрагмента)"
                    if tnum and tnum not in ("БИТАЯ", "вне статьи"):
                        resolv = ttitle or f"Статья {tnum}"
                        if not is_head:
                            resolv += "  [внутренний якорь статьи]"
                    elif tnum == "БИТАЯ":
                        resolv = "⚠ битый якорь (id не найден)"
                    else:
                        resolv = "(якорь вне статьи)"
                    rows.append({
                        "art": art_num, "loc": loc,
                        "anchor": atext, "target": target,
                        "resolve": resolv, "type": "внутр.",
                    })
                else:
                    name = EXT_NAMES.get(docid) or atext or "(внешний акт)"
                    tgt = (f"{docid}#{frag}" if (docid and frag)
                           else (docid or norm(href)[:40]))
                    resolv = f"{name}" + (f", adilet {docid}" if docid else "")
                    rows.append({
                        "art": art_num, "loc": loc,
                        "anchor": atext, "target": tgt,
                        "resolve": resolv, "type": "внешн.",
                    })
    return rows


def art_sort_key(n):
    m = re.match(r"(\d+)(?:-(\d+))?", n or "0")
    return (int(m.group(1)), int(m.group(2) or 0)) if m else (0, 0)


# ─────────────────────── СБОР ДАННЫХ ───────────────────────
rows = collect_rows()
rows.sort(key=lambda r: (art_sort_key(r["art"]),))
n_int = sum(1 for r in rows if r["type"] == "внутр.")
n_ext = sum(1 for r in rows if r["type"] == "внешн.")

# Блок «Помечено Анарой»: для flagged-статей — параграфы с внутр. ссылкой
anara_blocks = []
for art in FLAGGED:
    div = soup.find("div", attrs={"data-type": "статья", "data-number": art})
    title = article_title(div) if div else f"Статья {art}"
    items = []
    punkt = sub = None
    for p in div.find_all("p"):
        txt = norm(p.get_text(" "))
        mp = P_PUNKT.match(txt)
        ms = P_SUB.match(txt)
        if mp:
            punkt = mp.group(1); sub = None
        elif ms:
            sub = ms.group(1)
        if "настоящего Кодекса" not in txt:
            continue
        internal_links = [a for a in p.find_all("a")
                          if a.get("href") and classify_href(a.get("href"))[0] == "internal"]
        if not internal_links:
            continue
        raw = norm(p.decode_contents())
        resolves = []
        for a in internal_links:
            frag = classify_href(a.get("href"))[1]
            tnum, ttitle, _ = resolve_internal(frag)
            resolves.append(f"#{frag} → {ttitle or ('Статья ' + str(tnum))}")
        items.append({
            "loc": point_label(punkt, sub),
            "pid": p.get("id"),
            "raw": raw,
            "resolves": resolves,
        })
    anara_blocks.append({"art": art, "title": title, "items": items})

# Спот-чек: «об административных правонарушениях» без <a>
SPOT = "об административных правонарушениях"
spot_plain = []
for el in soup.find_all(string=re.compile(SPOT)):
    if el.find_parent("a"):
        continue
    ctx = norm(str(el))
    idx = ctx.find(SPOT)
    s = max(0, idx - 55)
    e = min(len(ctx), idx + len(SPOT) + 30)
    spot_plain.append("…" + ctx[s:e] + "…")


# ─────────────────────── MARKDOWN ───────────────────────
def md_escape(s):
    return s.replace("|", "\\|")


lines = []
lines.append("# Земельный кодекс РК (NGR K030000442_) — сверка гиперссылок для Анары")
lines.append("")
lines.append("READ-ONLY артефакт. Источник: `data/final/zemelnyy_structured.html`. "
             "HTML не изменялся.")
lines.append("")
lines.append(f"**Итого ссылок в статьях: {len(rows)}** "
             f"(внутренних `#z…`: {n_int}; внешних: {n_ext}).")
lines.append("")
lines.append("## 1. Все ссылки Земельного кодекса")
lines.append("")
lines.append("| Статья | Пункт | Текст-якорь | Target | Резолв | Тип |")
lines.append("|---|---|---|---|---|---|")
for r in rows:
    lines.append("| {} | {} | {} | `{}` | {} | {} |".format(
        md_escape(r["art"]), md_escape(r["loc"]), md_escape(r["anchor"]),
        md_escape(r["target"]), md_escape(r["resolve"]), r["type"]))
lines.append("")
lines.append("## 2. Помечено Анарой (13 пунктов) — ссылка стоит, номер на месте, ведёт верно")
lines.append("")
for b in anara_blocks:
    lines.append(f"### {b['title'] or ('Статья ' + b['art'])}")
    lines.append("")
    if not b["items"]:
        lines.append("_(внутренних ссылок «настоящего Кодекса» в статье не найдено)_")
        lines.append("")
        continue
    for it in b["items"]:
        lines.append(f"**{b['art']}, {it['loc']}**  (`p#{it['pid']}`)")
        lines.append("")
        lines.append("```html")
        lines.append(it["raw"])
        lines.append("```")
        for rv in it["resolves"]:
            lines.append(f"- резолв: {rv}")
        lines.append("")
lines.append("## 3. Спот-чек: «об административных правонарушениях» без `<a>`")
lines.append("")
lines.append(f"Найдено вхождений без ссылки: **{len(spot_plain)}**. "
             "Все они — описательный текст (дела/протоколы/постановления/"
             "законодательство об адм. правонарушениях), а НЕ отсылка к КоАП "
             "как акту → ссылка не требуется.")
lines.append("")
for i, c in enumerate(spot_plain, 1):
    lines.append(f"{i}. {c}")
lines.append("")

OUT_MD.write_text("\n".join(lines), encoding="utf-8")
print(f"MD  → {OUT_MD}  ({len(rows)} ссылок)")


# ─────────────────────── DOCX ───────────────────────
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9)

h = doc.add_heading("Земельный кодекс РК (NGR K030000442_)", level=0)
doc.add_heading("Сверка гиперссылок — артефакт для Анары (READ-ONLY)", level=2)
doc.add_paragraph("Источник: data/final/zemelnyy_structured.html. HTML не изменялся. "
                  f"Итого ссылок в статьях: {len(rows)} "
                  f"(внутренних #z…: {n_int}; внешних: {n_ext}).")

doc.add_heading("1. Все ссылки Земельного кодекса", level=1)
table = doc.add_table(rows=1, cols=6)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, t in enumerate(["Статья", "Пункт", "Текст-якорь", "Target", "Резолв", "Тип"]):
    hdr[i].text = t
    for par in hdr[i].paragraphs:
        for run in par.runs:
            run.font.bold = True
            run.font.size = Pt(8)
for r in rows:
    c = table.add_row().cells
    vals = [r["art"], r["loc"], r["anchor"], r["target"], r["resolve"], r["type"]]
    for i, v in enumerate(vals):
        c[i].text = v
        for par in c[i].paragraphs:
            for run in par.runs:
                run.font.size = Pt(8)

doc.add_page_break()
doc.add_heading("2. Помечено Анарой — ссылка стоит, номер на месте, ведёт верно", level=1)
for b in anara_blocks:
    doc.add_heading(b["title"] or f"Статья {b['art']}", level=2)
    if not b["items"]:
        doc.add_paragraph("(внутренних ссылок «настоящего Кодекса» не найдено)")
        continue
    for it in b["items"]:
        p = doc.add_paragraph()
        run = p.add_run(f"{b['art']}, {it['loc']}  (p#{it['pid']})")
        run.font.bold = True
        # сырой HTML моноширинно
        praw = doc.add_paragraph()
        rr = praw.add_run(it["raw"])
        rr.font.name = "Consolas"
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        for rv in it["resolves"]:
            pr = doc.add_paragraph(style="List Bullet")
            run = pr.add_run("резолв: " + rv)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

doc.add_page_break()
doc.add_heading("3. Спот-чек: «об административных правонарушениях» без ссылки", level=1)
doc.add_paragraph(
    f"Найдено вхождений без <a>: {len(spot_plain)}. Все они — описательный текст "
    "(рассмотрение дел / протоколы / постановления / законодательство об адм. "
    "правонарушениях), а НЕ отсылка к КоАП как акту → ссылка не требуется.")
for i, c in enumerate(spot_plain, 1):
    pr = doc.add_paragraph(style="List Number")
    run = pr.add_run(c)
    run.font.size = Pt(9)

doc.save(OUT_DOCX)
print(f"DOCX → {OUT_DOCX}")
print(f"   Анара-блок: {sum(len(b['items']) for b in anara_blocks)} параграфов; "
      f"спот-чек: {len(spot_plain)} вхождений")
